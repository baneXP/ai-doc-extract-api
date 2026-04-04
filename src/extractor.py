
# The Extractor function are Layer 1: get text of the files

#import libraries
import io #input-output
import pdfplumber # extract text from pdf
from docx import Document # used for .docx file 
from PIL import Image, ImageFilter, ImageEnhance

#PIL -> Pillow (image processsing libraries)
#Image -> open/load image
#ImageFilter -> filter for blur, sharpen
#ImageEnchance -> adjust contrast
import pytesseract #Lib for OCR
from pdf2image import convert_from_bytes #pdf -> image necessary for scaned image

def preprocess_image(image: Image.Image) -> Image.Image:
                # input type(PIL Image) return an image
    image=image.convert("L") #convert image to gray scale to increase contrast and accuracy
    
    image = ImageEnhance.Contrast(image).enhance(2.0) # create enhancer
    image = ImageEnhance.Sharpness(image).enhance(2.0)         
    #enchance contrast by 2x make text clearer
    image = image.filter(ImageFilter.MedianFilter())
    # sharpening to enchance edges
    return image

def extract_from_pdf(file_bytes: bytes) -> str:
                    #files in bytes     output in string
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        # io.BytesIO bytes -> pdf
        # pdfplumbe.open open pdf file
        # with -> auto close resources
            for page in pdf.pages:
                page_text = page.extract_text() #extract text from page
                if page_text: 
                    text += page_text + "\n"
    except Exception as e: #catch error in e
        print(f"[pdfplumber error] {e}")

    # scanned pdf (NO TEXT LAYER)
    if not text.strip(): #if no text extracted
        print("[INFO] No text layer found. Falling back to OCR....")
        try:
            images = convert_from_bytes(file_bytes, dpi=200) #pdf-> Image, dpi=200 higger ocr==better ocr
            for img in images:
                processed = preprocess_image(img) #processed image to increase the quality
                page_text = pytesseract.image_to_string(processed) #processed image to text
                text += page_text + "\n" #return final cleaned code.
        except Exception as e:
            print(f"[pdf2image/OCR error] {e}")

    return text.strip() #return final cleaned text

def extract_from_docx(file_bytes: bytes) -> str :
    
    text=""
    try:
        doc = Document(io.BytesIO(file_bytes)) #load word file from files
        
        #paragraphs
        for para in doc.paragraphs: #loop paragraphs
            if para.text.strip(): #ignores empty lines
                text += para.text + "\n"

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join( #join cells using "|" seperator
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text += row_text + "\n"

    except Exception as e:
        print(f"[docx extraction error] {e}")

    return text.strip()

def extract_from_image(file_bytes: bytes) -> str:
    text = ""
    try:
        image = Image.open(io.BytesIO(file_bytes))

        # upscale for better OCR
        width, height = image.size
        image = image.resize((width * 2, height * 2), Image.LANCZOS)

        # attempt 1: OCR on original color image
        # FIRST: use processed image (best accuracy)
        processed = preprocess_image(image)
        text = pytesseract.image_to_string(processed, config='--psm 6')

        # SECOND: try original image
        if len(text.strip()) < 30:
            text = pytesseract.image_to_string(image, config='--psm 6')

        # THIRD: try inverted image
        if len(text.strip()) < 30:
            from PIL import ImageOps
            inverted = ImageOps.invert(image.convert("RGB"))
            text = pytesseract.image_to_string(inverted, config='--psm 6')

    except Exception as e:
        print(f"[image OCR error] {e}")

    return text.strip()


def extract_text(file_bytes: bytes, file_type: str) -> str:
    # Dispatcherfunction decides which extrator to use

    file_type = file_type.lower().strip() #make text lower case and strip remove the spaces

    if file_type == "pdf":
        return extract_from_pdf(file_bytes)
    elif file_type == "docx":
        return extract_from_docx(file_bytes)
    elif file_type == "image":
        return extract_from_image(file_bytes) 
    else :
        raise ValueError(f"Unsupported file type: {file_type}")